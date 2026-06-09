import docx, pandas as pd
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
FONT='Times New Roman'
d=docx.Document(); d.styles['Normal'].font.name=FONT; d.styles['Normal'].font.size=Pt(11)
sec=d.sections[0]; sec.page_width=Inches(8.5); sec.page_height=Inches(11)
sec.top_margin=sec.bottom_margin=sec.left_margin=sec.right_margin=Inches(1)
def edge(par,name,val,sz='6'):
    el=OxmlElement('w:'+name); el.set(qn('w:val'),val); el.set(qn('w:sz'),sz); el.set(qn('w:space'),'0'); el.set(qn('w:color'),'auto'); par.append(el)
def three_line(table):
    table.style='Normal Table'; tblPr=table._tbl.tblPr
    b=OxmlElement('w:tblBorders'); edge(b,'top','single'); edge(b,'bottom','single')
    edge(b,'left','none'); edge(b,'right','none'); edge(b,'insideH','none'); edge(b,'insideV','none'); tblPr.append(b)
    for c in table.rows[0].cells:
        tcPr=c._tc.get_or_add_tcPr(); tcB=OxmlElement('w:tcBorders'); edge(tcB,'bottom','single'); tcPr.append(tcB)
def para(t,bold=False,it=False,sz=11):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=bold; r.italic=it; r.font.name=FONT; r.font.size=Pt(sz)
def block(df, cap, abbr, notes, fs=9):
    t=d.add_table(rows=1,cols=len(df.columns))
    for j,c in enumerate(df.columns):
        rr=t.rows[0].cells[j].paragraphs[0].add_run(str(c)); rr.bold=True; rr.font.name=FONT; rr.font.size=Pt(fs)
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,v in enumerate(row):
            rr=cells[j].paragraphs[0].add_run('' if pd.isna(v) else str(v)); rr.font.name=FONT; rr.font.size=Pt(fs)
    three_line(t)
    para(cap, bold=True); para('Abbreviations: '+abbr, it=False, sz=10); para('Notes: '+notes, sz=10); d.add_paragraph()
T='outputs/tables/'
para('SUPPLEMENTARY MATERIALS', bold=True, sz=13)
para('CS-MORT-6: Refining Cardiogenic Shock Risk Within SCAI Stages Through a Serially Computable Bedside Mortality Score'); d.add_paragraph()

# Supplementary tables are numbered in order of first citation in the main text (BMC requirement).

# S1 TRIPOD+AI checklist (first cited in Methods, reporting)
_tri=pd.DataFrame([
 {'Item':'1','Checklist item':'Title identifies model, population, outcome','Reported':'Yes','Section':'Title'},
 {'Item':'2','Checklist item':'Structured abstract','Reported':'Yes','Section':'Abstract'},
 {'Item':'3a-c','Checklist item':'Background, objectives, health inequalities','Reported':'Yes','Section':'Background'},
 {'Item':'4','Checklist item':'Study objectives','Reported':'Yes','Section':'Background'},
 {'Item':'5a-b','Checklist item':'Data sources and dates','Reported':'Yes','Section':'Methods'},
 {'Item':'6a-c','Checklist item':'Setting, eligibility, treatments','Reported':'Yes','Section':'Methods'},
 {'Item':'7','Checklist item':'Data pre-processing and quality','Reported':'Yes','Section':'Methods'},
 {'Item':'8','Checklist item':'Outcome definition and timing','Reported':'Yes','Section':'Methods'},
 {'Item':'9','Checklist item':'Predictor definitions and timing','Reported':'Yes','Section':'Methods, Table S2'},
 {'Item':'10','Checklist item':'Sample size / events per variable','Reported':'Yes','Section':'Methods'},
 {'Item':'11','Checklist item':'Missing data handling','Reported':'Yes','Section':'Methods, Table S9'},
 {'Item':'12a-g','Checklist item':'Analysis, model building, performance, recalibration','Reported':'Yes','Section':'Methods'},
 {'Item':'13','Checklist item':'Class imbalance','Reported':'Yes','Section':'Methods'},
 {'Item':'14','Checklist item':'Fairness approach','Reported':'Yes','Section':'Methods, Table S7'},
 {'Item':'15','Checklist item':'Model output and thresholds','Reported':'Yes','Section':'Methods, Table 2'},
 {'Item':'16','Checklist item':'Development vs evaluation differences','Reported':'Yes','Section':'Methods'},
 {'Item':'17-19','Checklist item':'Ethics, funding, conflicts, protocol, registration, PPI','Reported':'Yes','Section':'Declarations'},
 {'Item':'20a-c','Checklist item':'Participant flow, characteristics, distributions','Reported':'Yes','Section':'Figure 1, Table 1'},
 {'Item':'21','Checklist item':'Participants and events per analysis','Reported':'Yes','Section':'Results'},
 {'Item':'22','Checklist item':'Full model for third-party use','Reported':'Yes','Section':'Table 2, Table S5'},
 {'Item':'23a-b','Checklist item':'Performance with CIs, subgroups, clusters','Reported':'Yes','Section':'Results, Tables S7-S8'},
 {'Item':'24','Checklist item':'Model updating results','Reported':'Yes','Section':'Results'},
 {'Item':'25','Checklist item':'Interpretation including fairness','Reported':'Yes','Section':'Discussion'},
 {'Item':'26','Checklist item':'Limitations','Reported':'Yes','Section':'Discussion'},
 {'Item':'27a-c','Checklist item':'Implementation, user expertise, future research','Reported':'Yes','Section':'Discussion'}])
block(_tri,'Table S1. TRIPOD+AI Reporting Checklist',
 'CI, confidence interval; PPI, patient and public involvement; TRIPOD+AI, Transparent Reporting of a multivariable prediction model for Individual Prognosis Or Diagnosis, artificial intelligence extension.',
 'Reporting follows the TRIPOD+AI statement (Collins et al., BMJ 2024). The table gives the location of each item; items that do not apply to a retrospective secondary-data analysis are marked as not applicable rather than reported as met.')

# S2 variable definitions (first cited in Methods, candidate predictors)
block(pd.DataFrame([
 {'Variable':'Lactate','Definition':'Most recent lactate','Units':'mmol/L','Time window':'≤24 h'},
 {'Variable':'Urine output','Definition':'First-24h urine / weight / observed hours','Units':'mL/kg/h','Time window':'0-24 h'},
 {'Variable':'OHCA','Definition':'Cardiac-arrest diagnosis with emergency admission','Units':'Binary','Time window':'At admission'},
 {'Variable':'Age','Definition':'Age at admission','Units':'years','Time window':'At admission'},
 {'Variable':'Blood urea nitrogen','Definition':'Most recent BUN','Units':'mg/dL','Time window':'≤24 h'},
 {'Variable':'Red cell distribution width','Definition':'Most recent RDW','Units':'%','Time window':'≤24 h'},
 {'Variable':'Anion gap','Definition':'Sodium − (chloride + bicarbonate), computed','Units':'mmol/L','Time window':'≤24 h'}]),
 'Table S2. Variable Definitions and Measurement Windows',
 'BUN, blood urea nitrogen; OHCA, out-of-hospital cardiac arrest; RDW, red cell distribution width.',
 'All laboratory values represent the most recent measurement within the window. The anion gap is computed identically in both databases so that thresholds transport.')

# S3 model-class comparison (first cited in Methods, model development)
block(pd.DataFrame([
 {'Model':'Elastic-net logistic (selected)','AUROC':'0.792','Calibration slope':'0.93'},
 {'Model':'LASSO logistic','AUROC':'0.792','Calibration slope':'0.95'},
 {'Model':'Random forest','AUROC':'0.790','Calibration slope':'1.61'},
 {'Model':'Gradient boosting','AUROC':'0.796','Calibration slope':'0.75'}]),
 'Table S3. Model-Class Comparison',
 'AUROC, area under the receiver operating characteristic curve; LASSO, least absolute shrinkage and selection operator.',
 'Cross-validated discrimination was similar across model classes (gradient boosting 0.796, elastic-net 0.792); the tree-ensemble models showed less stable out-of-the-box calibration (slopes 1.61 and 0.75), and penalized logistic regression was chosen for interpretability and the transparent integer score it supports rather than for superior discrimination.')

# S4 cut-point sensitivity + note classifier (first cited in Methods, integer score derivation)
block(pd.DataFrame([
 {'Analysis':'Integer cut-points: guideline/data-derived','Result':'AUROC 0.763'},
 {'Analysis':'Integer cut-points: distribution quintiles','Result':'AUROC 0.765'},
 {'Analysis':'Integer cut-points: round-number clinical','Result':'AUROC 0.766'},
 {'Analysis':'Note classifier, note-level PPV (all affirmed, n=1,199)','Result':'96.5%'},
 {'Analysis':'Note classifier, note-only inclusion-driving (n=282)','Result':'89.7%'}]),
 'Table S4. Cut-Point Sensitivity and Note-Classifier Validation',
 'AUROC, area under the receiver operating characteristic curve; PPV, positive predictive value.',
 'Integer-score discrimination was robust across three independent cut-point schemes, confirming the cut-points are not arbitrary. The note-classifier positive predictive value reflects precision among affirmed mentions; recall was not estimated, so complete case ascertainment is not claimed, and a code-only phenotype was retained as a sensitivity analysis.')

# S5 continuous model specification (first cited in Methods, integer score derivation)
block(pd.read_csv(T+'S10_coefficients.csv'),
 'Table S5. Continuous Model Specification for Third-Party Use',
 'BUN, blood urea nitrogen; OHCA, out-of-hospital cardiac arrest; RDW, red cell distribution width; SD, standard deviation.',
 'Full specification of the continuous models (TRIPOD+AI item 22). Each predictor is winsorized to the stated 1st to 99th percentile range; missing values are imputed with the median; values are then standardized as z = (value minus mean) divided by SD. The linear predictor equals the intercept plus the sum of beta times z, and the predicted probability equals 1 / (1 + exp(minus linear predictor)). Standardized intercepts are -0.499 for the lactate development model and -0.575 for the anion-gap deployable model; odds ratios per standard deviation equal exp(beta). The anion-gap variant substitutes the harmonized anion gap for lactate, and its coefficients differ because the anion gap captures part of the renal and acid-base signal otherwise carried by blood urea nitrogen. Models were penalized logistic regression (L2 penalty, C=0.5, random seed 42).')

# S6 sensitivity analyses (first cited in Methods, statistical analysis)
block(pd.read_csv(T+'S_sensitivity_analyses.csv').rename(columns={'analysis':'Analysis','n':'N','mortality':'Mortality','continuous_AUROC':'Continuous AUROC','integer_AUROC':'Integer AUROC'}),
 'Table S6. Sensitivity Analyses Across Cohort Definitions',
 'AUROC, area under the receiver operating characteristic curve; CS, cardiogenic shock; OHCA, out-of-hospital cardiac arrest.',
 'The primary cohort retains mixed septic-cardiogenic shock because culture results are unavailable at the time of scoring (excluding on them would introduce look-ahead bias). Discrimination is higher in pure cardiogenic shock and lower within the culture-positive subgroup, consistent with the score capturing cardiac rather than infective mortality, although the higher values may partly reflect the more homogeneous case mix after exclusion. All AUROC values are 5-fold cross-validated.')

# S7 subgroup (fairness) performance (first cited in Methods, statistical analysis)
block(pd.DataFrame([
 {'Subgroup':'Male','N':1858,'Mortality':'36.4%','AUROC (95% CI)':'0.785 (0.762-0.807)','Slope':1.04,'CITL':'-0.072'},
 {'Subgroup':'Female','N':1245,'Mortality':'41.1%','AUROC (95% CI)':'0.767 (0.738-0.792)','Slope':0.90,'CITL':'+0.097'},
 {'Subgroup':'White','N':1935,'Mortality':'36.8%','AUROC (95% CI)':'0.770 (0.748-0.791)','Slope':1.01,'CITL':'-0.050'},
 {'Subgroup':'Black','N':317,'Mortality':'36.6%','AUROC (95% CI)':'0.780 (0.725-0.831)','Slope':0.85,'CITL':'-0.324'},
 {'Subgroup':'Other/Unknown','N':693,'Mortality':'44.9%','AUROC (95% CI)':'0.799 (0.763-0.829)','Slope':1.03,'CITL':'+0.335'}]),
 'Table S7. Subgroup (Fairness) Performance',
 'AUROC, area under the receiver operating characteristic curve; CI, confidence interval; CITL, calibration-in-the-large.',
 'Discrimination was similar across sex and race subgroups, with overlapping confidence intervals (TRIPOD+AI item 14). Calibration was well centred in the larger subgroups; the wider calibration-in-the-large in the Black (n=317) and Other or Unknown subgroups reflects limited sample size rather than established miscalibration, and subgroup-specific recalibration is advisable before subgroup-level use. Hispanic (n=84) and Asian (n=74) groups were too small for stable estimates.')

# S8 cluster heterogeneity (first cited in Methods, statistical analysis)
block(pd.DataFrame([{'Hospitals (≥25 CS patients)':24,'Median AUROC':0.764,'IQR':'0.713-0.791','Range':'0.621-0.957','Pooled eICU AUROC':0.748}]),
 'Table S8. Cluster Heterogeneity Across eICU Hospitals',
 'AUROC, area under the receiver operating characteristic curve; CS, cardiogenic shock; IQR, interquartile range.',
 'Median per-hospital discrimination matched the pooled estimate; the wide range (0.621 to 0.957) reflects small per-hospital samples, as few as 25 patients, with correspondingly wide per-site confidence intervals rather than demonstrated site-level failure (TRIPOD+AI items 12d and 23b).')

# S9 missing data analysis and multiple imputation (first cited in Methods, statistical analysis)
block(pd.DataFrame([
 {'Variable':'Lactate','Missing':'19.4%','Mortality if missing':'28.7%','Mortality if measured':'40.6%','Pattern':'MNAR (less sick missing)'},
 {'Variable':'Urine output','Missing':'8.3%','Mortality if missing':'56.2%','Mortality if measured':'36.7%','Pattern':'MNAR (sicker missing)'}]),
 'Table S9. Missing Data Analysis and Multiple Imputation',
 'AUROC, area under the receiver operating characteristic curve; MICE, multiple imputation by chained equations; MNAR, missing not at random.',
 'Discrimination stratified by lactate status: measured AUROC 0.789, missing (imputed) 0.701, with preserved calibration in both (slope ~1.0). Multiple imputation AUROC 0.775 versus median imputation 0.778, confirming robustness. The anion gap substitutes for missing lactate.')

# S10 diagnostic accuracy (first cited in Methods, statistical analysis)
block(pd.read_csv(T+'S11_diagnostic.csv'),
 'Table S10. Diagnostic Accuracy at Integer Score Thresholds',
 'LR+, positive likelihood ratio; LR−, negative likelihood ratio; NPV, negative predictive value; PPV, positive predictive value.',
 'A score of 8 or more identifies a high-risk group suitable for rule-in (high specificity and positive likelihood ratio). Because mortality remained substantial even at the lowest scores, the score is not suitable for ruling out death.')

# S11 eICU baseline (first cited in Results, cohort characteristics)
block(pd.read_csv(T+'S1_eicu_baseline.csv'),
 'Table S11. eICU External Validation Cohort Baseline Characteristics',
 'eICU, eICU Collaborative Research Database; OHCA, out-of-hospital cardiac arrest.',
 'Values are median (interquartile range) for continuous variables and n (%) for categorical variables, measured within 24 hours. P-values compare survivors and non-survivors (Mann-Whitney U or chi-square).')

# S12 risk categories (first cited in Results, discrimination and calibration)
_s2=pd.read_csv(T+'T4_risk_categories.csv').rename(columns={'category':'Risk category','score':'Score','MIMIC':'MIMIC-IV mortality','eICU_lactate':'eICU (lactate)','eICU_AG':'eICU (anion gap)'})
block(_s2,'Table S12. Risk Stratification by Score Category',
 'eICU, eICU Collaborative Research Database; MIMIC-IV, Medical Information Mart for Intensive Care IV.',
 'Categories are presented as relative risk strata. The same frozen score cut-points were applied to both cohorts; low-risk mortality was 13.7% in MIMIC-IV versus 10.7% (lactate) and 14.1% (anion gap) in eICU, a difference of 0.4 to 3.0 percentage points, in contrast to the marked drift of the predecessor score.')

# S13 head-to-head comparison (first cited in Results, Table 3 footnote)
block(pd.DataFrame([
 {'Score':'CS-MORT-6 (anion gap)','Scorable':'50% all six observed / 95% AG term / 100% imputed','AUROC (95% CI)':'0.748 (0.724-0.771)','Calibration / ECE':'slope 0.95 / 0.046'},
 {'Score':'CS-MORT-6 (lactate)','Scorable':'25% all six observed / 49% lactate term / 100% imputed','AUROC (95% CI)':'0.757 (0.733-0.779)','Calibration / ECE':'slope 0.85 / —'},
 {'Score':'BOS,MA2','Scorable':'60% complete checklist','AUROC (95% CI)':'0.743 (0.713-0.770)','Calibration / ECE':'— / 0.146'},
 {'Score':'CS-MORT-8 (predecessor)','Scorable':'100% imputed','AUROC (95% CI)':'0.745 (0.720-0.769)','Calibration / ECE':'— / —'}]),
 'Table S13. Head-to-Head Comparison With Existing Scores in eICU',
 'AUROC, area under the receiver operating characteristic curve; CI, confidence interval; ECE, expected calibration error.',
 'DeLong test on commonly scorable patients: CS-MORT-6 (anion gap) versus BOS,MA2, difference 0.005, p=0.749 (equivalent). CardShock and IABP-SHOCK II were not computable in eICU (ejection fraction and post-procedural coronary flow unavailable). Applied off the shelf, BOS,MA2 had a larger expected calibration error (0.146 versus 0.046), but this gap is almost entirely calibration-in-the-large driven by the different base rate; after recalibration to the external base rate the residual calibration error was comparable (0.007 for BOS,MA2 versus 0.019 for CS-MORT-6). CS-MORT-6 transported with only a minor intercept adjustment.')

# (Calibration and decision-curve figures promoted to the main manuscript as Figures 2 and 3;
#  the supplement now contains tables only.)
d.save('manuscript/CS_MORT_6_SUPPLEMENTARY.docx')
print('wrote manuscript/CS_MORT_6_SUPPLEMENTARY.docx with', len(d.tables),'tables')
