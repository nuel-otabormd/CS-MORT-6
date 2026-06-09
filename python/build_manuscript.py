import re, docx, pandas as pd
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
FONT='Times New Roman'
d=docx.Document(); st=d.styles['Normal']; st.font.name=FONT; st.font.size=Pt(11)
# Page layout: US Letter, standard 1-inch margins (fixes wide/odd default borders)
_sec=d.sections[0]; _sec.page_width=Inches(8.5); _sec.page_height=Inches(11)
_sec.top_margin=_sec.bottom_margin=_sec.left_margin=_sec.right_margin=Inches(1)
def TABLE(df, fs=10):
    t=d.add_table(rows=1, cols=len(df.columns)); t.style='Table Grid'
    t.autofit=True
    for j,c in enumerate(df.columns):
        cell=t.rows[0].cells[j]; r=cell.paragraphs[0].add_run(str(c)); r.bold=True; r.font.name=FONT; r.font.size=Pt(fs)
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,v in enumerate(row):
            r=cells[j].paragraphs[0].add_run('' if pd.isna(v) else str(v)); r.font.name=FONT; r.font.size=Pt(fs)
    d.add_paragraph()
CIT=re.compile(r'\{\{([\d,–\-\s]+)\}\}')
def add_runs(p,text):
    pos=0
    for m in CIT.finditer(text):
        if m.start()>pos:
            r=p.add_run(text[pos:m.start()]); r.font.name=FONT; r.font.size=Pt(11)
        sup=p.add_run(m.group(1)); sup.font.superscript=True; sup.font.name=FONT; sup.font.size=Pt(11)
        pos=m.end()
    if pos<len(text):
        r=p.add_run(text[pos:]); r.font.name=FONT; r.font.size=Pt(11)
def H1(t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.name=FONT; r.font.size=Pt(11)
def H2(t):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.name=FONT; r.font.size=Pt(11)
def P(t,sp=2.0,bold=False,center=False):
    p=d.add_paragraph(); p.paragraph_format.line_spacing=sp
    if center: p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    elif sp>=1.5 and not bold: p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold:
        r=p.add_run(t); r.bold=True; r.font.name=FONT; r.font.size=Pt(11)
    else: add_runs(p,t)
    return p

def SUPER(segments, sp=1.0):
    p=d.add_paragraph(); p.paragraph_format.line_spacing=sp
    for text, sup in segments:
        r=p.add_run(text); r.font.name=FONT; r.font.size=Pt(11); r.font.superscript=sup
    return p

# ============ TITLE PAGE ============
P('CS-MORT-6: REFINING CARDIOGENIC SHOCK RISK WITHIN SCAI STAGES THROUGH A SERIALLY COMPUTABLE BEDSIDE MORTALITY SCORE',sp=1.0,bold=True)
P('Running Title: Serial Bedside Risk Score in Cardiogenic Shock',sp=1.0)
SUPER([('Emmanuel Otabor, MD',False),('a',True),('; Kevin Bryan Lo, MD',False),('b',True),('; Ayoyimika Okunlola, MBBS',False),('c',True),('; Justin Lam, MD',False),('a',True),('; Laith Alomari, MD',False),('a',True),('; Michael Hamilton, MD',False),('a',True),('; Abiodun Idowu, MD',False),('a',True),('; Abdulraheem Hassan, MD',False),('d',True),('; Olayinka Afolabi-Brown, MD',False),('a',True)])
SUPER([('a',True),(' Department of Medicine, Jefferson Einstein Philadelphia Hospital, Philadelphia, PA, USA',False)])
SUPER([('b',True),(' Harvard Medical School, Division of Cardiovascular Medicine, Brigham and Women’s Hospital, Boston, MA, USA',False)])
SUPER([('c',True),(' Nottingham University Hospital NHS Trust, Nottingham, UK',False)])
SUPER([('d',True),(' Saint Peter’s University Hospital, New Brunswick, NJ, USA',False)])
P('Conflicts of Interest Statement: All authors declare that they have no known competing financial interests or relationships.',sp=1.0)
P('Corresponding Author:',sp=1.0)
P('Emmanuel Otabor, MD (ORCID: 0009-0006-7902-6909)',sp=1.0)
P('Department of Medicine, Jefferson Einstein Philadelphia Hospital, PA, USA.',sp=1.0)
P('Email: emmanuel.otabor@jefferson.edu',sp=1.0)
P('Manuscript details:',sp=1.0,bold=True)
P('Number of references: 35',sp=1.0)
P('Number of tables: 3',sp=1.0)
P('Number of figures: 3',sp=1.0)

# ============ ABSTRACT (250 words) ============
d.add_page_break()
H1('ABSTRACT')
P('Background: Cardiogenic shock carries an in-hospital mortality of 30-50%. Existing risk scores require echocardiography or post-procedural data, are computed once, and do not refine the qualitative Society for Cardiovascular Angiography and Interventions (SCAI) shock stage. We sought to develop and externally validate a serially computable bedside score and to test whether it resolves mortality risk within SCAI stages.')
P('Methods: Using MIMIC-IV (2008-2019), we identified 3,103 adults with documented cardiogenic shock across all intensive care units. Penalized logistic regression with bootstrap stability selection produced a six-variable score (lactate, urine output, out-of-hospital cardiac arrest, age, blood urea nitrogen, and red cell distribution width), converted to integer points by the Sullivan method. The frozen model was externally validated in the eICU database (n=1,866), with anion gap substituted for lactate where unavailable.')
P('Results: In-hospital mortality was 38.3% (MIMIC-IV) and 34.5% (eICU). Cross-validated discrimination was an area under the curve of 0.778 internally and 0.748-0.757 externally, with strong calibration (slope 0.95). The integer score ranged from 0-15 and was not overfit on bootstrap correction, and discrimination matched the best externally computable score. Within identical SCAI stages, where stages C and D had nearly identical mean mortality, the score separated low-risk and high-risk tertiles by 24-53 percentage points, and a worsening 24-48 hour trajectory identified deterioration (58% versus 25% mortality from the same baseline).')
P('Conclusions: CS-MORT-6 is a serially computable bedside score that refines SCAI staging and tracks deterioration, with transportable calibration and within-stage risk resolution at comparable discrimination.')
P('Keywords: cardiogenic shock; risk score; mortality prediction; SCAI stage; critical care',sp=1.0)

# ============ CLINICAL PERSPECTIVE ============
H1('CLINICAL PERSPECTIVE')
H2('What Is New?')
P('CS-MORT-6 is a six-variable bedside score for in-hospital mortality in cardiogenic shock that requires neither echocardiography nor post-procedural data and is computable serially throughout the intensive care stay.')
P('Within a single SCAI shock stage, the score separated low-risk and high-risk patients by 24-53 percentage points in two independent databases, and a worsening score over 24-48 hours identified clinical deterioration that the baseline value could not.')
H2('What Are the Clinical Implications?')
P('The score provides a quantitative risk estimate that complements the qualitative SCAI stage and can be repeated as the patient evolves, supporting decisions about escalation of care, mechanical circulatory support evaluation, and goals-of-care discussions.')
P('Because the anion gap can substitute for lactate where lactate is unavailable, the score remains computable across heterogeneous intensive care settings.')

# ============ INTRODUCTION ============
d.add_page_break()
H1('Introduction')
P('Cardiogenic shock remains among the most lethal conditions in cardiovascular medicine, with in-hospital mortality of 30-50% despite contemporary advances in revascularization, mechanical circulatory support, and systems of care.{{1,2}} Incidence has risen over the past two decades, and outcomes remain poor across the spectrum of ischemic and non-ischemic causes.{{3,4}}')
P('Accurate, early risk stratification is central to the management of cardiogenic shock, informing the intensity of monitoring, the timing of escalation and mechanical circulatory support evaluation, candidacy for transfer within regionalized networks, and goals-of-care discussions.{{2,7}} The SCAI shock classification has been widely adopted to communicate severity and consistently tracks mortality across its five stages.{{5,6,7}} The stage, however, is a qualitative ordinal category, and patients within the same stage exhibit substantial residual heterogeneity in outcome that the classification was not designed to resolve.')
P('Several quantitative scores exist but each carries a practical constraint. The CardShock score incorporates left ventricular ejection fraction,{{8}} and the IABP-SHOCK II score requires the post-percutaneous-intervention coronary flow grade,{{9}} so neither is computable at the bedside before imaging or catheterization, nor in databases lacking those inputs. Echo-free checklist scores avoid these requirements but rely on support-dependent vital signs and are applied once at a fixed time point.{{10}} No existing tool is designed for serial use as the patient evolves, and none has been shown to refine the SCAI stage.')
P('We therefore sought to develop and externally validate a bedside mortality score for cardiogenic shock that uses only routinely available variables, that can be recomputed at any time during the intensive care stay, and that quantifies risk within SCAI stages. We prespecified that the aim was deployability and stage refinement rather than a discrimination record, and we report performance honestly against the best externally computable comparator.')

# ============ METHODS ============
H1('Methods')
H2('Data Availability')
P('This study used MIMIC-IV (version 3.1) for development and the eICU Collaborative Research Database for external validation, both available through PhysioNet upon completion of credentialing and data use agreements.{{11,12,13}} Direct data sharing is not permitted, but qualified researchers can reproduce the analysis from the original databases. The complete analytical code, cohort queries, and the prespecified analysis plan are publicly available on GitHub and archived at Zenodo.')
H2('Study Design and Data Sources')
P('This retrospective cohort study used MIMIC-IV for model development and internal validation and the eICU database for external validation. MIMIC-IV contains intensive care records from the Beth Israel Deaconess Medical Center (2008-2019),{{12}} whereas eICU comprises more than 200,000 admissions from 208 hospitals across the United States (2014-2015).{{13}} Both databases carry pre-existing institutional review board approval, and no additional ethical approval was required.{{12,13}} Reporting followed the TRIPOD+AI statement, and risk of bias was self-assessed with PROBAST.{{29,30}}')
H2('Study Population')
P('Adults aged 18 years or older with cardiogenic shock were identified across all intensive care units using a documentation-anchored phenotype. A qualifying admission required documentation of cardiogenic shock, defined as an International Classification of Diseases code (ICD-10 R57.0 or ICD-9 785.51) or an affirmative mention of cardiogenic shock in the discharge summary, together with at least one objective criterion within the first 24 hours, namely a systolic blood pressure below 90 mmHg or a mean arterial pressure below 65 mmHg, a lactate of at least 2 mmol/L, or a vasoactive, inotropic, or mechanical circulatory support requirement. Discharge summaries were classified with a context-aware natural language pipeline based on the ConText algorithm, which separates affirmed mentions from negated, historical, hypothetical, and uncertain contexts;{{25}} on manual review of affirmed mentions the note-level positive predictive value was 96.5%; because this reflects precision among affirmed mentions rather than recall, a code-only definition was also examined as a sensitivity analysis and complete case ascertainment is not claimed. One index admission was retained per patient, and no minimum length of stay was imposed, so that early deaths were retained and immortal time bias avoided. Equivalent criteria identified the external cohort, in which cardiogenic shock was ascertained from the Acute Physiology and Chronic Health Evaluation diagnosis string; a harmonized cohort additionally requiring an objective criterion was examined as a sensitivity analysis. Patient flow is illustrated in Figure 1.')
P('Figure 1. Study Flow Diagram',sp=1.0,bold=True)
P('Study flow diagram illustrating patient selection from the MIMIC-IV and eICU databases. Abbreviations: ICU, intensive care unit; ICD, International Classification of Diseases; SCAI, Society for Cardiovascular Angiography and Interventions.',sp=1.0)
H2('Outcome')
P('The primary outcome was in-hospital mortality during the index hospitalization. The secondary outcome was 30-day mortality. In-hospital mortality was selected for its relevance to acute decision-making and its objective ascertainment, and because 87.8% of 30-day deaths in this cohort occurred during the index hospitalization, indicating close concordance between the endpoints.')
H2('Candidate Predictors')
P('Candidate predictors were extracted as the most recent value recorded up to the 24-hour reference time rather than the worst value, so that the score reflects the current physiological state and can improve as the patient improves;{{15}} all measurements were restricted to times at or before the reference time to preclude the use of future information. The final model comprised lactate, urine output rate, out-of-hospital cardiac arrest, age, blood urea nitrogen, and red cell distribution width. Blood urea nitrogen was retained on the basis of ADHERE registry data demonstrating its independent prognostic value in acute heart failure,{{24}} and red cell distribution width on the basis of its established prognostic association in heart failure.{{23}} The urine output rate was normalized by body weight and by the actual observed hours, capped at 24. For external deployment the anion gap, computed identically in both databases from sodium, chloride, and bicarbonate,{{19}} substituted for lactate where lactate was unavailable. Complete variable definitions are provided in Table S2.')
H2('Missing Data')
P('Both perfusion variables were missing not at random and in opposite directions: patients without a recorded lactate had lower mortality, reflecting selective ordering, whereas patients without a recorded urine output had higher mortality. Median imputation served as the explicit point-of-care rule, with the harmonized anion gap as a prespecified substitution for missing lactate. Discrimination and calibration were reported separately for patients with and without a measured lactate, and multiple imputation by chained equations was examined as a sensitivity analysis (Table S2).')
H2('Model Development')
P('Continuous variables were winsorized at the first and 99th percentiles and standardized with parameters fitted on the development data. The modelling algorithm was selected by comparing penalized logistic regression, random forests, and gradient boosting; cross-validated discrimination was similar across algorithms, within 0.01 of an area under the curve, and the tree-ensemble models showed less stable out-of-the-box calibration, consistent with evidence that flexible methods offer no advantage over logistic regression for structured clinical prediction,{{16}} penalized logistic regression was adopted for its interpretability and the transparent integer score it enables, rather than on the basis of a discrimination advantage. Predictors were selected by bootstrap stability selection, after which treatment-dependent, computed, near-outcome, and composite features were removed by prespecified clinical rules; the six retained variables were selected in every resample. With 1,188 events for six parameters, the cohort provided approximately 198 events per parameter, exceeding accepted minimums.{{17}}')
H2('Integer Score Derivation')
P('A bedside integer score was derived by the method of Sullivan and colleagues,{{18}} categorizing each variable, fitting an ordinal logistic model, and scaling coefficients to integer points. Categorical thresholds were anchored to established cut-points where available, including lactate thresholds from cardiogenic shock and lactate-clearance literature{{21,22}} and urine output thresholds from KDIGO acute kidney injury criteria,{{20}} and to the observed dose-response otherwise; integer discrimination was robust across alternative threshold schemes. Risk categories were defined to provide clinical anchoring, adequate distribution, and a monotonic mortality increase in both cohorts, and were presented as relative strata subject to local recalibration.')
H2('Statistical Analysis')
P('Discrimination was summarized by the area under the receiver operating characteristic curve with confidence intervals from 1,000 bootstrap iterations, and internal optimism by 500-resample bootstrap correction and five-fold cross-validation; a single split-sample approach was avoided as statistically inefficient.{{15}} Calibration was assessed by the calibration slope, calibration-in-the-large, a flexible calibration curve, and the Brier score,{{26}} and clinical utility by decision curve analysis.{{27}} For external validation the entire pipeline was frozen on the development data and applied unchanged to eICU. Head-to-head comparison was performed against the echo-free score of Yamga and colleagues{{10}} and the original eight-variable predecessor, with differences tested by the DeLong method;{{28}} the CardShock and IABP-SHOCK II scores were not externally computable because ejection fraction and post-procedural coronary flow are unavailable in eICU.{{8,9}} The SCAI stage{{6}} was operationalized using a drug, device, and arrest escalation hierarchy, and within-stage stratification assessed by score tertiles. Serial behaviour was evaluated by recomputing the frozen integer score at 24 and 48 hours and relating the change to mortality among 48-hour survivors. Prespecified sensitivity analyses examined code-confirmed cardiogenic shock, sepsis-excluded and culture-confirmed-infection-excluded cohorts,{{14}} a comfort-care-excluded cohort, an arrest-removed model, complete-case and multiple-imputation analyses, subgroup performance by sex and race, and heterogeneity across eICU hospitals. Analyses used Python 3.9, figures were generated in R, and a fixed random seed was applied throughout.')

# ============ RESULTS ============
H1('Results')
H2('Cohort Characteristics')
P('The development cohort comprised 3,103 adults with documented cardiogenic shock, with an in-hospital mortality of 38.3% and a 30-day mortality of 41.6%. Mechanical circulatory support was used in 628 patients (20.2%). The external cohort comprised 1,866 admissions with an in-hospital mortality of 34.5%. Baseline characteristics are presented in Table 1 and, for the external cohort, in Table S11.')
P('Table 1. Baseline Characteristics of the MIMIC-IV Derivation Cohort',sp=1.0,bold=True)
TABLE(pd.read_csv('outputs/tables/T1_baseline.csv'))
P('Values are median (interquartile range) for continuous variables and n (%) for categorical variables, measured within 24 hours. Abbreviations: AMI, acute myocardial infarction; SCAI, Society for Cardiovascular Angiography and Interventions.',sp=1.0)
H2('Model and Integer Score')
P('The final model comprised six variables. In the integer formulation, lactate contributed up to 4 points, out-of-hospital cardiac arrest 3 points, and age, blood urea nitrogen, red cell distribution width, and urine output up to 2 points each, for a total range of 0-15 (Table 2). Discrimination exceeded that of lactate alone by 0.10 in the area under the curve, confirming that the additional variables contributed meaningfully.')
P('Table 2. CS-MORT-6 Integer Scoring System',sp=1.0,bold=True)
TABLE(pd.DataFrame([
 {'Variable':'Lactate (mmol/L)','Category':'<2 / 2-4 / ≥4','Points':'0 / 2 / 4'},
 {'Variable':'Urine output (mL/kg/h)','Category':'≥1 / 0.5-1 / <0.5','Points':'0 / 1 / 2'},
 {'Variable':'Out-of-hospital cardiac arrest','Category':'No / Yes','Points':'0 / 3'},
 {'Variable':'Age (years)','Category':'<65 / 65-80 / ≥80','Points':'0 / 1 / 2'},
 {'Variable':'Blood urea nitrogen (mg/dL)','Category':'<25 / 25-45 / ≥45','Points':'0 / 1 / 2'},
 {'Variable':'Red cell distribution width (%)','Category':'<14.5 / 14.5-16 / ≥16','Points':'0 / 1 / 2'}]))
P('Total score ranges from 0-15. Where lactate is unavailable, the anion gap (computed from sodium, chloride, and bicarbonate) substitutes, scored <12 / 12-18 / ≥18 as 0 / 2 / 4.',sp=1.0)
H2('Discrimination and Calibration')
P('Cross-validated discrimination was an area under the curve of 0.778 internally for the continuous model and 0.765 for the integer score, with negligible bootstrap optimism (0.766 corrected). On external validation the frozen model achieved an area under the curve of 0.757 for the lactate formulation and 0.748 for the deployable anion-gap formulation, with a calibration slope of 0.95 and calibration-in-the-large near zero (Table 3; flexible calibration curves are shown in Figure 2). Internal calibration was excellent (slope 0.99). Risk categories rose monotonically and showed a concordant gradient between cohorts, with the low-risk category at 13.7% in MIMIC-IV and 10.7% to 14.1% in eICU, in contrast to the marked low-risk drift of the predecessor (Table S12).')
P('Table 3. Model Performance and External Comparison',sp=1.0,bold=True)
_t3=pd.read_csv('outputs/tables/T3_performance.csv')
TABLE(_t3)
P('Discrimination is the area under the receiver operating characteristic curve (AUROC); MIMIC values are 5-fold cross-validated, eICU values are from the frozen model. For external comparison, CS-MORT-6 (anion gap) achieved an AUROC of 0.748 versus 0.743 for the BOS,MA2 score on commonly scorable patients (DeLong p=0.749, statistically equivalent). The score uses only routinely available variables with no echocardiographic or invasive inputs; the anion gap substituting for lactate was observed in approximately 95% of patients, and all six predictors were jointly observed in approximately 50%, comparable to a complete BOS,MA2 checklist in approximately 60%, with median imputation applied otherwise. Risk-category mortality and the full comparison are provided in Tables S12 and S13. Abbreviations: AG, anion gap; CITL, calibration-in-the-large.',sp=1.0)
H2('Comparison With Existing Scores')
P('On the patients scorable for both, CS-MORT-6 and the echo-free comparator were statistically indistinguishable in discrimination (area under the curve 0.748 versus 0.743; DeLong p=0.749), as were comparisons with the predecessor. In contrast to the CardShock and IABP-SHOCK II scores, which require an ejection fraction and post-procedural coronary flow and were not computable in eICU, CS-MORT-6 uses only routinely available variables and can be recomputed serially. The anion gap substituting for lactate was observed in approximately 95% of patients, and all six predictors were jointly observed in approximately 50%, comparable to a complete BOS,MA2 checklist in approximately 60%, with median imputation applied otherwise. The deployable formulation provided net benefit across the clinically relevant range of decision thresholds (Figure 3).')
H2('Within-Stage Resolution and Serial Behaviour')
P('Across SCAI stages, mortality rose as expected, yet stages C and D had nearly identical mean mortality, confirming the limited within-spectrum separation of the ordinal stage. Within each stage, the score separated low-risk and high-risk tertiles by 24-53 percentage points, and this resolution replicated in eICU (Figure 4). Among patients alive at 48 hours, recomputing the score improved discrimination over the stale 24-hour value, and among patients with an identical 24-hour score those whose score worsened by 48 hours had a mortality of 57.8% compared with 24.5% for those who improved (Figure 5).')
P('Figure 4. Within-SCAI-Stage Risk Resolution',sp=1.0,bold=True)
P('In-hospital mortality by CS-MORT-6 score tertile (low, intermediate, and high) within each SCAI shock stage, shown separately for the MIMIC-IV derivation cohort and the eICU external validation cohort. Within every stage the score separates a low-risk from a high-risk group, demonstrating quantitative refinement of the qualitative stage. Abbreviations: SCAI, Society for Cardiovascular Angiography and Interventions.',sp=1.0)
P('Figure 5. Serial Score Trajectory and Clinical Deterioration',sp=1.0,bold=True)
P('In-hospital mortality among patients alive at 48 hours, according to the change in CS-MORT-6 score between 24 and 48 hours (improved, stable, or worsened), shown for all 48-hour survivors (left panel) and restricted to patients with an identical intermediate 24-hour score (right panel). A worsening score identifies clinical deterioration that is not apparent from the baseline value.',sp=1.0)
H2('Sensitivity, Fairness, and Heterogeneity')
P('Discrimination was stable across cohort definitions and was higher in pure cardiogenic shock, reaching 0.806 when sepsis was excluded and 0.819 when culture-confirmed infection was excluded, although these higher values may partly reflect the more homogeneous case mix after exclusion rather than a difference in mechanism. Discrimination was similar across sex and race subgroups, with overlapping confidence intervals, and across eICU hospitals, although calibration in the smaller race subgroups was imprecise and per-hospital estimates varied widely owing to small samples (Tables S6 to S8 and S11). In a complete-case analysis restricted to patients with all six predictors observed, discrimination and calibration were maintained rather than degraded (deployable anion-gap formulation AUROC 0.760, 95% CI 0.726-0.796, n=928, calibration slope 0.95; lactate formulation 0.780, 95% CI 0.740-0.821, n=469, calibration slope 0.77), confirming that performance was not an artifact of imputation. Consistent with missingness not at random in opposite directions, the complete-case subsets differed in baseline mortality, which was lower where urine output was observed and higher where lactate was observed, so these estimates reflect a selected case mix and are reported as a robustness check rather than a preferred estimate.')

# ============ DISCUSSION ============
H1('Discussion')
P('We developed and externally validated CS-MORT-6, a six-variable bedside score for in-hospital mortality in cardiogenic shock that can be computed serially throughout the intensive care stay and that refines the SCAI shock stage. The score discriminated well internally and retained good discrimination with strong calibration when frozen and applied across 208 hospitals. Its discrimination matched that of the best externally computable echo-free score and of the eight-variable predecessor, and we interpret this parity directly rather than obscure it: the value of CS-MORT-6 lies not in superior rank-ordering, where echo-free scores in cardiogenic shock have plateaued near an area under the curve of 0.74-0.78, but in computability without echocardiographic or invasive inputs, calibration that transports, serial applicability, and the resolution of risk within SCAI stages.')
P('Two features distinguish the score from existing tools. First, it quantifies risk within SCAI stages. The classification communicates severity and tracks mortality across its five stages,{{5,6,7}} yet within a single stage the residual heterogeneity in outcome is large; in our data stages C and D carried nearly identical mean mortality despite representing different points on the management pathway. By separating low-risk from high-risk patients by 24-53 percentage points within each stage, in two independent databases and under more than one operationalization of the stage, CS-MORT-6 provides quantitative refinement that complements the qualitative classification rather than competing with it. Second, the score is designed for serial use. Existing cardiogenic shock scores are computed once at a fixed time,{{8,9,10}} whereas recomputing CS-MORT-6 at 48 hours improved on the stale 24-hour value, and the direction of change carried independent prognostic information, with patients whose score worsened experiencing more than twice the mortality of those whose score improved from the same baseline. This behaviour aligns with the established observation that lactate clearance predicts survival more reliably than any single lactate measurement in cardiogenic shock.{{21,22}}')
P('The variables comprising the score map to distinct pathophysiological domains. Arterial lactate, the most heavily weighted variable, reflects tissue hypoperfusion and anaerobic metabolism and is a consistent prognostic marker in cardiogenic shock.{{21,22}} Urine output and blood urea nitrogen capture renal hypoperfusion and the cardiorenal interaction that characterizes advancing shock,{{24}} red cell distribution width reflects chronic systemic illness and maladaptive erythropoiesis and is prognostic across heart failure populations,{{23}} and out-of-hospital cardiac arrest identifies a superimposed anoxic insult that strongly influences survival. Computing the anion gap directly from electrolytes, rather than relying on the institution-specific laboratory value, allowed its thresholds to transport between databases,{{19}} which both improved external calibration and provided a principled substitution for lactate where lactate was unavailable, addressing a limitation that has constrained prior bedside scores.')
P('Our findings should be read against the tools they are intended to complement. The CardShock score discriminates well but incorporates left ventricular ejection fraction,{{8}} and the IABP-SHOCK II score requires the post-percutaneous-intervention coronary flow grade,{{9}} so neither could be computed in the external database and neither is available at the bedside before imaging or catheterization. The larger echo-free score reported higher discrimination in its development cohort{{10}} but relies on support-dependent vital signs and, applied off the shelf to our higher-mortality external population, substantially under-predicted observed risk, whereas the predicted risks of CS-MORT-6 transported with only a minor intercept adjustment. The decision to retain mixed septic-cardiogenic shock in the development cohort, rather than exclude it on the basis of cultures unavailable at the time of scoring, preserves the intended deployment population and avoids conditioning on future information; that discrimination was higher still in pure cardiogenic shock is consistent with the score capturing cardiac rather than infective risk, although higher discrimination after exclusion may also reflect a more homogeneous case mix.')
P('These properties carry practical implications. Positioned as a stratification and monitoring trigger rather than a determinant of individual treatment, a high or rising score may prompt earlier evaluation for mechanical circulatory support, transfer within a regionalized hub-and-spoke network, or a goals-of-care discussion, whereas a low and stable score may support de-escalation of monitoring.{{2,33}} Because the score is recomputed from routinely available values, it is suited to the serial reassessment that contemporary shock-team protocols emphasize, and its quantitative output may sharpen communication that is currently anchored on the qualitative SCAI stage. Whether embedding the score within a clinical decision-support workflow improves patient outcomes will require prospective evaluation.{{35}}')
P('This study has limitations. Development was single-centre, although external validation spanned 208 hospitals. Discrimination is moderate and statistically indistinguishable from the best available comparator, so the score should inform rather than dictate decisions. The documentation-anchored phenotype relies in part on discharge-era coding and notes that are not available in real time, so the score applies to recognized cardiogenic shock; the dynamic signal rises on the three time-varying variables while the remaining three are static; and the SCAI stage in the external database was approximated with a coarse operationalization constrained by the available data. Finally, the calibration advantage over the checklist comparator reflects in part the similar baseline mortality of the development and validation cohorts, and local recalibration remains advisable when the score is applied to populations of differing severity.')

# ============ CONCLUSION ============
H1('Conclusion')
P('CS-MORT-6 is a six-variable bedside score for in-hospital mortality in cardiogenic shock that is computable serially from routinely available data, refines mortality risk within SCAI stages, and tracks clinical deterioration. Externally validated across a large multicentre cohort with strong calibration and broad applicability, it offers a practical complement to qualitative shock staging that may support escalation, mechanical circulatory support evaluation, and goals-of-care decisions. Prospective evaluation and impact assessment are warranted.')

# ============ REFERENCES ============
H1('References')
refs=[
 'Samsky MD, Morrow DA, Proudfoot AG, Hochman JS, Thiele H, Rao SV. Cardiogenic Shock After Acute Myocardial Infarction: A Review. JAMA. 2021;326(18):1840-1850.',
 'van Diepen S, Katz JN, Albert NM, et al. Contemporary Management of Cardiogenic Shock: A Scientific Statement From the American Heart Association. Circulation. 2017;136(16):e232-e268.',
 'Osman M, Syed M, Patibandla S, et al. Fifteen-Year Trends in Incidence of Cardiogenic Shock Hospitalization and In-Hospital Mortality in the United States. J Am Heart Assoc. 2021;10(15):e021061.',
 'Berg DD, Bohula EA, Morrow DA. Epidemiology and Causes of Cardiogenic Shock. Curr Opin Crit Care. 2021;27(4):401-408.',
 'Baran DA, Grines CL, Bailey S, et al. SCAI Clinical Expert Consensus Statement on the Classification of Cardiogenic Shock. Catheter Cardiovasc Interv. 2019;94(1):29-37.',
 'Naidu SS, Baran DA, Jentzer JC, et al. SCAI SHOCK Stage Classification Expert Consensus Update: A Review and Incorporation of Validation Studies. J Am Coll Cardiol. 2022;79(9):933-946.',
 'Jentzer JC, van Diepen S, Barsness GW, et al. Cardiogenic Shock Classification to Predict Mortality in the Cardiac Intensive Care Unit. J Am Coll Cardiol. 2019;74(17):2117-2128.',
 'Harjola VP, Lassus J, Sionis A, et al. Clinical picture and risk prediction of short-term mortality in cardiogenic shock. Eur J Heart Fail. 2015;17(5):501-509.',
 'Pöss J, Köster J, Fuernau G, et al. Risk Stratification for Patients in Cardiogenic Shock After Acute Myocardial Infarction. J Am Coll Cardiol. 2017;69(15):1913-1920.',
 'Yamga E, Mantena S, Rosen D, et al. Optimized Risk Score to Predict Mortality in Patients With Cardiogenic Shock in the Cardiac Intensive Care Unit. J Am Heart Assoc. 2023;12(13):e029232.',
 'Goldberger AL, Amaral LAN, Glass L, et al. PhysioBank, PhysioToolkit, and PhysioNet. Circulation. 2000;101(23):e215-e220.',
 'Johnson AEW, Bulgarelli L, Shen L, et al. MIMIC-IV, a Freely Accessible Electronic Health Record Dataset. Sci Data. 2023;10:1.',
 'Pollard TJ, Johnson AEW, Raffa JD, et al. The eICU Collaborative Research Database. Sci Data. 2018;5:180178.',
 'Singer M, Deutschman CS, Seymour CW, et al. The Third International Consensus Definitions for Sepsis and Septic Shock (Sepsis-3). JAMA. 2016;315(8):801-810.',
 'Steyerberg EW, Vergouwe Y. Towards Better Clinical Prediction Models: Seven Steps for Development and an ABCD for Validation. Eur Heart J. 2014;35(29):1925-1931.',
 'Christodoulou E, Ma J, Collins GS, Steyerberg EW, Verbakel JY, Van Calster B. A systematic review shows no performance benefit of machine learning over logistic regression for clinical prediction models. J Clin Epidemiol. 2019;110:12-22.',
 'Riley RD, Snell KIE, Ensor J, et al. Minimum sample size for developing a multivariable prediction model: PART II, binary and time-to-event outcomes. Stat Med. 2019;38(7):1276-1296.',
 'Sullivan LM, Massaro JM, D’Agostino RB Sr. Presentation of Multivariate Data for Clinical Use: The Framingham Study Risk Score Functions. Stat Med. 2004;23(10):1631-1660.',
 'Kraut JA, Madias NE. Serum anion gap: its uses and limitations in clinical medicine. Clin J Am Soc Nephrol. 2007;2(1):162-174.',
 'Kidney Disease: Improving Global Outcomes (KDIGO) Acute Kidney Injury Work Group. KDIGO Clinical Practice Guideline for Acute Kidney Injury. Kidney Int Suppl. 2012;2(1):1-138.',
 'Fuernau G, Desch S, de Waha-Thiele S, et al. Arterial Lactate in Cardiogenic Shock: Prognostic Value of Clearance Versus Single Values. JACC Cardiovasc Interv. 2020;13(19):2208-2216.',
 'Marbach JA, Stone S, Schwartz B, et al. Lactate Clearance Is Associated With Improved Survival in Cardiogenic Shock: A Systematic Review and Meta-Analysis. J Card Fail. 2021;27(10):1082-1089.',
 'Felker GM, Allen LA, Pocock SJ, et al. Red Cell Distribution Width as a Novel Prognostic Marker in Heart Failure: Data From the CHARM Program and the Duke Databank. J Am Coll Cardiol. 2007;50(1):40-47.',
 'Fonarow GC, Adams KF Jr, Abraham WT, et al. Risk Stratification for In-Hospital Mortality in Acutely Decompensated Heart Failure: Classification and Regression Tree Analysis. JAMA. 2005;293(5):572-580.',
 'Harkema H, Dowling JN, Thornblade T, Chapman WW. ConText: An algorithm for determining negation, experiencer, and temporal status from clinical reports. J Biomed Inform. 2009;42(5):839-851.',
 'Van Calster B, McLernon DJ, van Smeden M, et al. Calibration: the Achilles heel of predictive analytics. BMC Med. 2019;17(1):230.',
 'Vickers AJ, Elkin EB. Decision Curve Analysis: A Novel Method for Evaluating Prediction Models. Med Decis Making. 2006;26(6):565-574.',
 'DeLong ER, DeLong DM, Clarke-Pearson DL. Comparing the Areas under Two or More Correlated Receiver Operating Characteristic Curves. Biometrics. 1988;44(3):837-845.',
 'Collins GS, Moons KGM, Dhiman P, et al. TRIPOD+AI statement: updated guidance for reporting clinical prediction models that use regression or machine learning methods. BMJ. 2024;385:e078378.',
 'Moons KGM, Wolff RF, Riley RD, et al. PROBAST: A Tool to Assess Risk of Bias and Applicability of Prediction Model Studies: Explanation and Elaboration. Ann Intern Med. 2019;170(1):W1-W33.',
 'Hochman JS, Sleeper LA, Webb JG, et al. Early Revascularization in Acute Myocardial Infarction Complicated by Cardiogenic Shock. N Engl J Med. 1999;341(9):625-634.',
 'Thiele H, Zeymer U, Neumann FJ, et al. Intraaortic Balloon Support for Myocardial Infarction with Cardiogenic Shock. N Engl J Med. 2012;367(14):1287-1296.',
 'Tehrani BN, Truesdell AG, Psotka MA, et al. A Standardized and Comprehensive Approach to the Management of Cardiogenic Shock. JACC Heart Fail. 2020;8(11):879-891.',
 'Heidenreich PA, Bozkurt B, Aguilar D, et al. 2022 AHA/ACC/HFSA Guideline for the Management of Heart Failure. J Am Coll Cardiol. 2022;79(17):e263-e421.',
 'Sutton RT, Pincock D, Baumgart DC, et al. An Overview of Clinical Decision Support Systems. npj Digit Med. 2020;3:17.',
]
for i,r in enumerate(refs,1):
    p=d.add_paragraph(); p.paragraph_format.line_spacing=1.5; run=p.add_run(f"{i}. {r}"); run.font.name=FONT; run.font.size=Pt(11)

d.save('manuscript/CS_MORT_6_FULL_MANUSCRIPT.docx')
# word count of body
import docx as _d
doc=_d.Document('manuscript/CS_MORT_6_FULL_MANUSCRIPT.docx')
words=sum(len(p.text.split()) for p in doc.paragraphs)
print(f"wrote manuscript/CS_MORT_6_FULL_MANUSCRIPT.docx | total words ~{words} | refs {len(refs)}")
# abstract word count
ab=[p.text for p in doc.paragraphs]
abstract=" ".join(t for t in ab if t.startswith(('Background:','Methods:','Results:','Conclusions:')))
print("Abstract words:", len(abstract.split()))
