"""Generate the CS-MORT-6 manuscript-preparation document package.
All docs use a CONSISTENT font: Times New Roman, 12 pt (journal-standard, adjustable)."""
import docx
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
FONT='Times New Roman'; SIZE=12
def newdoc():
    d=docx.Document(); s=d.styles['Normal']; s.font.name=FONT; s.font.size=Pt(SIZE)
    return d
def H(d,t,lvl=1):
    p=d.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.name=FONT; r.font.size=Pt(14 if lvl==1 else 12)
    return p
def P(d,t,bold=False):
    p=d.add_paragraph(); r=p.add_run(t); r.font.name=FONT; r.font.size=Pt(SIZE); r.bold=bold
    return p
def save(d,name): d.save(f'manuscript/{name}'); print('wrote manuscript/'+name)

# ---------- 00 README ----------
d=newdoc(); H(d,'CS-MORT-6 Manuscript Package — Overview')
P(d,'This folder contains the preparation documents for the CS-MORT-6 manuscript. All documents use Times New Roman 12 pt for consistency (adjustable on request). Working title: "Refining Cardiogenic Shock Risk Within SCAI Stages: Development and External Validation of a Serially Computable Bedside Mortality Score." Reporting standard: TRIPOD+AI (Collins et al., BMJ 2024). Target: JAHA / Circulation: Cardiovascular Quality and Outcomes / Critical Care / European Journal of Heart Failure.')
P(d,'Contents:',True)
for line in ['01_References_and_PDF_needs — full reference list (reused from CS-MORT-8 plus new), and exactly which PDFs are still required.',
 '02_Sepsis3_in_CS_methodological_note — why Sepsis-3 over-identifies sepsis in cardiogenic shock, the culture-confirmed refinement, and sensitivity results.',
 '03_Reusable_content_from_CSMORT8 — which Introduction, Methods, Discussion text and references carry over and what must change.',
 '04_Figure_and_image_specifications — high-quality figure standards (R, >=600 DPI / vector) and the planned figure list.',
 '05_Manuscript_skeleton — section-by-section skeleton mapped to TRIPOD+AI items and the two-contribution story.']:
    P(d,'• '+line)
save(d,'00_README.docx')

# ---------- 01 References & PDF needs ----------
d=newdoc(); H(d,'Reference List and PDF Requirements')
P(d,'The CS-MORT-8 manuscript carried 35 references; most transfer directly. Below: (A) references reused as-is, (B) references needing the PDF for content/exact definitions, (C) NEW references the rebuild requires, and (D) the prioritized list of PDFs to obtain.')
H(d,'A. Reused references (citation sufficient, no PDF needed)',2)
for r in ['Samsky 2021 JAMA (CS after AMI review) — Intro epidemiology','van Diepen 2017 Circulation (AHA CS statement) — Intro','Osman 2021 JAHA (CS trends) — Intro','Berg 2021 Curr Opin Crit Care (CS epidemiology) — Intro','Tehrani 2020 JACC HF; Tehrani 2019 JACC; Tehrani 2022 JACC HF (CS systems of care) — Intro/Discussion','Hochman 1999 NEJM SHOCK; Thiele 2012 NEJM IABP-SHOCK II — Intro/Discussion','Johnson 2023 Sci Data (MIMIC-IV); Pollard 2018 Sci Data (eICU); PhysioNet — Methods','Steyerberg 2014 EHJ; Van Calster (calibration) — Methods','Vickers 2006 (DCA); DeLong 1988 — Methods','KDIGO 2012 (AKI/urine output thresholds) — Methods','Sutton 2020 npj Digit Med (CDSS); Heidenreich 2022 HF guideline — Discussion']:
    P(d,'• '+r)
H(d,'B. Reused but PDF needed (extract exact variable definitions / numbers)',2)
for r in ['Naidu 2022 JACC (SCAI SHOCK update) — exact stage definitions for Leg 1','Baran 2019 Catheter Cardiovasc Interv (original SCAI consensus) — staging','Jentzer 2019 JACC (SCAI mortality in CICU) — within-stage / SCAI validation context','Harjola 2015 Eur J Heart Fail (CardShock) — exact comparator variables','Pöss 2017 JACC (IABP-SHOCK II) — exact comparator variables','Fuernau 2020 JACC Cardiovasc Interv (lactate clearance in CS) — dynamic/lactate story','Marbach 2020 (lactate clearance meta-analysis) — lactate rationale','Fonarow 2005 JAMA (ADHERE; BUN as top predictor) — BUN rationale','Sullivan 2004 Stat Med (integer score method) — Methods integer derivation','Yamga 2023 JAHA (BOS,MA2) — HAVE THIS PDF already (comparator definitions confirmed)']:
    P(d,'• '+r)
H(d,'C. NEW references required for the rebuild',2)
for r in ['Collins 2024 BMJ (TRIPOD+AI statement) — reporting standard; checklist already obtained','Christodoulou 2019 J Clin Epidemiol (ML vs logistic regression) — model-class bake-off justification','Riley 2020 BMJ/Stat Med (minimum sample size for prediction models) — events-per-variable justification','Kraut & Madias 2007 NEJM (serum anion gap) — anion-gap harmonization rationale','Singer 2016 JAMA (Sepsis-3 definition) — sepsis sensitivity + the over-call discussion','Felker 2007 JACC (RDW in heart failure, CHARM) — RDW predictor rationale [VERIFY exact cite]','Harkema 2009 J Biomed Inform (ConText algorithm) — note phenotyping','Eyre 2021 (medspaCy) — note phenotyping implementation [VERIFY]','Wolff 2019 Ann Intern Med (PROBAST) — risk-of-bias self-assessment']:
    P(d,'• '+r)
H(d,'D. PDFs to obtain, prioritized (paywalled — please download)',2)
P(d,'HIGH priority (content needed, likely paywalled): Naidu 2022 JACC; Jentzer 2019 JACC; Harjola 2015 EJHF; Poss 2017 JACC; Fuernau 2020 JACC Cardiovasc Interv; Fonarow 2005 JAMA (ADHERE); Felker 2007 JACC (RDW); Marbach 2020.')
P(d,'OPEN ACCESS (I can retrieve): Collins 2024 BMJ (TRIPOD+AI); Johnson 2023 / Pollard 2018 (MIMIC/eICU); Christodoulou 2019; PROBAST; Singer 2016 (JAMA, may be free); Kraut & Madias 2007.')
P(d,'ALREADY HAVE: Yamga 2023 (BOS,MA2).')
P(d,'Note: per anti-fabrication policy, no DOI or finding will be cited without the verified source; items marked [VERIFY] are confirmed against PubMed before writing.')
save(d,'01_References_and_PDF_needs.docx')

# ---------- 02 Sepsis-3 note ----------
d=newdoc(); H(d,'Sepsis-3 in Cardiogenic Shock: Limitations and a Culture-Confirmed Refinement')
P(d,'Rationale for scrutiny.',True)
P(d,'Sepsis-3 (Singer et al., JAMA 2016) defines sepsis as a suspected or documented infection accompanied by an acute rise in the Sequential Organ Failure Assessment (SOFA) score of at least two points. Both components are problematic in cardiogenic shock.')
P(d,'Problem 1 — SOFA is non-specific in shock.',True)
P(d,'The SOFA score awards points for hypotension and vasopressor use (cardiovascular), elevated creatinine and low urine output (renal), hyperbilirubinemia (hepatic), and thrombocytopenia (coagulation). In cardiogenic shock these derangements arise from systemic hypoperfusion, not infection, so a SOFA increase of two or more points is nearly universal regardless of whether infection is present. The SOFA arm therefore does little to distinguish infected from uninfected patients in this population.')
P(d,'Problem 2 — suspicion of infection is triggered by the routine shock work-up.',True)
P(d,'The MIMIC-IV operationalization of suspected infection requires an antibiotic administration paired with a body-fluid culture within a defined window. Patients presenting in undifferentiated shock routinely receive empiric antibiotics and have cultures drawn while infection is being excluded, which satisfies the suspicion criterion even when no infection is ultimately found.')
P(d,'Empirical demonstration in our cohort.',True)
P(d,'Sepsis-3 flagged 48.4% of the cardiogenic-shock cohort. However, among these Sepsis-3-positive patients, only 48.6% had a positive culture; the remainder were flagged on suspicion (empiric antibiotics plus cultures drawn) together with a non-specific SOFA elevation, with negative cultures. Sepsis-3 therefore over-identifies sepsis in cardiogenic shock.')
P(d,'A more specific definition and the sensitivity result.',True)
P(d,'Restricting to culture-confirmed infection (positive culture from the suspicion-of-infection derived table) yields a more specific phenotype of genuine mixed septic-cardiogenic shock. CS-MORT-6 discrimination is robust and, if anything, higher in cohorts that exclude infection: full cohort cross-validated AUROC 0.778; Sepsis-3-excluded 0.806; culture-confirmed-infection-excluded 0.819. The score is therefore not an artifact of sepsis mortality.')
P(d,'Recommendation for the manuscript.',True)
P(d,'Report Sepsis-3 prevalence with an explicit statement of its non-specificity in cardiogenic shock, and present a culture-confirmed-infection sensitivity cohort as the more specific definition of mixed shock. Retain mixed shock within the primary cohort (it is clinically real and excluding on the basis of a future or empiric infection work-up would introduce bias), and use the culture-confirmed analysis to demonstrate robustness. Data source: physionet-data.mimiciv_3_1_derived.suspicion_of_infection (positive_culture field).')
save(d,'02_Sepsis3_in_CS_methodological_note.docx')

# ---------- 03 Reusable content ----------
d=newdoc(); H(d,'Reusable Content from the CS-MORT-8 Manuscript')
P(d,'The prior manuscript (FINAL CS-MORT-8_MANUSCRIPT.docx) provides a strong template. The following maps what transfers and what must change for CS-MORT-6.')
H(d,'Introduction — largely reusable',2)
P(d,'Reuse: the epidemiology paragraph (40,000-50,000 AMI-CS cases annually in the United States; 30-50% mortality), the risk-stratification rationale (timing of intervention, mechanical-support candidacy, goals-of-care), and the survey of existing scores (CardShock requires ejection fraction; IABP-SHOCK II requires post-PCI TIMI flow). CHANGE: the objective paragraph must be rewritten around the two contributions (refining SCAI staging and serial computability), not "another bedside score," and must state up front that discrimination is on par with existing tools. ADD a sentence on known health inequalities in cardiogenic shock care (TRIPOD+AI item 3c).')
H(d,'Methods — partially reusable',2)
P(d,'Reuse: data-source descriptions (MIMIC-IV, eICU), the bedside-variable philosophy, integer-score (Sullivan) method, calibration/DCA/DeLong descriptions, the external-validation framing. CHANGE/ADD: the all-ICU documentation-anchored cohort (replaces CCU-only), the ConText note phenotype and its validation, most-recent (not worst-value) features, the anion-gap harmonization, the missing-data MNAR treatment, the SCAI operationalization, the dynamic 24-48h analysis, the fairness and cluster-heterogeneity analyses, and the Sepsis-3/culture sensitivity.')
H(d,'Discussion — partially reusable',2)
P(d,'Reuse: the per-variable pathophysiology paragraph (lactate as hypoperfusion marker; renal and hepatic congestion), the comparison-to-existing-tools structure, and the clinical-application-by-risk-category paragraph. CHANGE: lead with the two contributions and discrimination parity (honest framing), fold in the calibration/transportability nuance (the anion-gap harmonization), add the within-SCAI-stage and serial-deterioration interpretation, and add a fairness paragraph (TRIPOD+AI item 25). Drop bilirubin/hemoglobin-specific text (those variables are not in CS-MORT-6).')
H(d,'References — ~30 of 35 transfer',2)
P(d,'See 01_References_and_PDF_needs. Drop: Lundberg SHAP (no SHAP in the rebuild) unless interpretability is retained; Pencina reclassification (not used); Cherbi hemoglobin (variable dropped) unless cited as CS-cohort context. Replace TRIPOD 2015 (Collins 2015) with TRIPOD+AI 2024. Add the new methodological references.')
save(d,'03_Reusable_content_from_CSMORT8.docx')

# ---------- 04 Figure specs ----------
d=newdoc(); H(d,'Figure and Image Specifications (High Quality)')
P(d,'All figures are generated in R (ggplot2 and related) and exported at publication quality. Standards:')
for s in ['Resolution: minimum 300 DPI; target 600 DPI for raster, or vector (PDF/EPS) for line-based figures (calibration, DCA, forest, ROC).',
 'Format: TIFF (lossless, preferred by AHA/Circulation journals) or high-resolution PNG; PDF vector for line art.',
 'Dimensions set explicitly: single column 3.5 in, double column 7.0 in; height to suit.',
 'Fonts in figures: sans-serif (Arial/Helvetica), minimum 8 pt after scaling.',
 'Color: colorblind-safe palettes (viridis or RColorBrewer Set2/Paired); interpretable in grayscale.',
 'Consistent theme across all figures via a shared theme_publication.R.']:
    P(d,'• '+s)
H(d,'Planned main figures',2)
for s in ['Figure 1 — CONSORT participant flow (MIMIC development and eICU validation).',
 'Figure 2 — Calibration (loess) internal and external (anion-gap variant), with slope/CITL/Brier.',
 'Figure 3 — Decision-curve analysis (MIMIC; eICU vs BOS,MA2 with recalibrated comparator).',
 'Figure 4 — Within-SCAI-stage risk resolution panel (MIMIC and eICU): Contribution 1.',
 'Figure 5 — Score trajectory and deterioration (24 to 48 h strata): Contribution 2.']:
    P(d,'• '+s)
P(d,'Supplementary figures: integer-score calibration; ROC curves with confidence bands; subgroup (fairness) performance forest plot; per-hospital AUROC distribution (cluster heterogeneity). Current draft figures are in outputs/figures/ at 300 DPI and will be re-exported at final resolution.')
save(d,'04_Figure_and_image_specifications.docx')

# ---------- 05 Manuscript skeleton ----------
d=newdoc(); H(d,'CS-MORT-6 Manuscript Skeleton (TRIPOD+AI-aligned)')
P(d,'Title: Refining Cardiogenic Shock Risk Within SCAI Stages: Development and External Validation of a Serially Computable Bedside Mortality Score. Running title: Serial bedside risk score in cardiogenic shock.')
secs=[('Abstract','Structured; TRIPOD+AI-for-Abstracts items. State development (MIMIC, n=3,103) and external validation (eICU, n=1,866); discrimination ~0.78 internal / 0.75 external; within-SCAI-stage resolution and serial deterioration as the contributions; discrimination parity with existing scores stated honestly.'),
 ('Introduction','Epidemiology (reuse); risk-stratification need; gap 1 (existing scores need echo/angiography or are one-shot) and gap 2 (SCAI is qualitative/coarse); objective = two contributions + external validation; health-inequality sentence (item 3c).'),
 ('Methods','Data sources and dates (5a,5b); setting/centres (6a); cohort definition and rationale (6b) with ConText phenotype + validation; outcome (8a) in-hospital primary, 30-day secondary; predictors most-recent <=24h leak-safe (9a,9b); sample size/EPV (10); missing data MNAR + median + AG substitution + MI (11); model class via bake-off + penalized LR (12c); predictor handling incl. anion-gap harmonization (12b); internal bootstrap + CV (12a); performance measures (12e); recalibration (12f); cluster-heterogeneity method (12d); class-imbalance none and why (13); fairness approach (14); model output + integer card + thresholds (15); development-vs-validation differences (16); ethics/funding/COI/protocol/registration/PPI (17-19); SCAI operationalization; dynamic 24-48h; sensitivity suite.'),
 ('Results','Participant flow + Table 1 (20a,20b); development-vs-validation distributions (20c); N and events per analysis (21); discrimination + calibration internal/external with CIs (23a) incl. fairness subgroups; cluster heterogeneity (23b); full model for reuse (22); comparators + DeLong + applicability (23a); recalibration (24); Contribution 1 within-SCAI-stage (Figure 4); Contribution 2 serial trajectory (Figure 5); sensitivity analyses (phenotype/Sepsis-3/culture, withdrawal, OHCA, complete-case, MI, note-classifier PPV).'),
 ('Discussion','Principal findings (two contributions + parity-with-honesty) (25, incl. fairness); per-variable pathophysiology (reuse); comparison with existing tools; calibration/transportability nuance; clinical applications by risk category and serial use; limitations (26); implementation guidance — handling missing/poor input and required user expertise (27a,27b); future directions (27c).'),
 ('Conclusion','A serially computable, six-variable bedside score that refines SCAI staging and tracks deterioration, externally validated, with deployability and calibration advantages at comparable discrimination.'),
 ('Other','Data availability (PhysioNet credentialing), code availability (GitHub + Zenodo DOI), ethics statement, funding, conflicts, author contributions, TRIPOD+AI checklist as supplement, PROBAST self-assessment.')]
for t,b in secs:
    H(d,t,2); P(d,b)
save(d,'05_Manuscript_skeleton.docx')
print('\nDONE: 6 documents in manuscript/ (Times New Roman 12 pt)')
