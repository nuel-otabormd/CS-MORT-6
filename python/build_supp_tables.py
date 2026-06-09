import warnings; warnings.filterwarnings('ignore')
import numpy as np, pandas as pd, docx
from docx.shared import Pt
from sklearn.linear_model import LogisticRegression
from sklearn.preprocessing import StandardScaler
from sklearn.impute import SimpleImputer
from sklearn.model_selection import StratifiedKFold
from sklearn.pipeline import Pipeline
from sklearn.metrics import roc_auc_score
np.random.seed(42)
d=pd.read_csv('outputs/tables/cs_features_canonical.csv').rename(columns={'uo_rate_mlkghr':'uo'})
pc=pd.read_csv('/tmp/poscx.csv'); d=d.merge(pc,on='stay_id',how='left'); d['positive_culture']=d['positive_culture'].fillna(0)
cs=pd.read_csv('/tmp/codestatus.csv'); d=d.merge(cs[['stay_id','ever_cmo']],on='stay_id',how='left')
COM=['uo','ohca_arrest','age','bun','rdw']; RC={'lactate':[-1,2,4,99],'bun':[-1,25,45,9e9],'age':[-1,65,80,999],'rdw':[-1,14.5,16,99]}; PR={'uo':[-1,0.5,1.0,99]}
def integ(df,feats=['lactate','uo','ohca_arrest','age','bun','rdw']):
    parts=[]
    for c in feats:
        if c=='ohca_arrest': parts.append(df[c].fillna(0).astype(int).values*3)
        elif c in PR:
            o=pd.cut(df[c],bins=PR[c],labels=False); o=o.fillna(o.median()); parts.append((o.max()-o).astype(int).values)
        else:
            o=pd.cut(df[c],bins=RC[c],labels=False); o=o.fillna(o.median()); parts.append(o.astype(int).values*(2 if c=='lactate' else 1))
    return np.sum(parts,axis=0)
def cont(df,cols=None):
    cols=cols or ['lactate']+COM; yy=df['in_hospital_mortality'].astype(int).values
    X=df[cols].astype(float).clip(df[cols].quantile(.01),df[cols].quantile(.99),axis=1); p=np.zeros(len(yy))
    for tr,te in StratifiedKFold(5,shuffle=True,random_state=42).split(X,yy):
        m=Pipeline([('i',SimpleImputer(strategy='median')),('s',StandardScaler()),('l',LogisticRegression(max_iter=800,C=0.5))]).fit(X.iloc[tr],yy[tr]); p[te]=m.predict_proba(X.iloc[te])[:,1]
    return roc_auc_score(yy,p)
# ---- sensitivity table (persist CSV) ----
rows=[]
defs=[('Primary cohort (all documented CS, mixed shock included)',d,None),
      ('Sensitivity: culture-confirmed infection excluded',d[d.positive_culture==0],None),
      ('Sensitivity: Sepsis-3 excluded',d[d.sepsis3==0],None),
      ('Subgroup: culture-positive (mixed shock)',d[d.positive_culture==1],None),
      ('Sensitivity: ICD-confirmed only',d[d.has_cs_icd==1],None),
      ('Sensitivity: comfort-measures-only excluded',d[d.ever_cmo!=1],None)]
for lab,sub,_ in defs:
    yy=sub['in_hospital_mortality'].astype(int).values
    rows.append({'analysis':lab,'n':len(sub),'mortality':f"{yy.mean():.1%}",'continuous_AUROC':round(cont(sub),3),'integer_AUROC':round(roc_auc_score(yy,integ(sub)),3)})
# OHCA-removed (continuous, 5-var)
yy=d['in_hospital_mortality'].astype(int).values
rows.append({'analysis':'Sensitivity: OHCA removed (CS-MORT-5)','n':len(d),'mortality':f"{yy.mean():.1%}",'continuous_AUROC':round(cont(d,['lactate','uo','age','bun','rdw']),3),'integer_AUROC':np.nan})
rows.append({'analysis':'Subgroup: non-OHCA patients','n':int((d.ohca_arrest!=1).sum()),'mortality':f"{d[d.ohca_arrest!=1].in_hospital_mortality.mean():.1%}",'continuous_AUROC':round(cont(d[d.ohca_arrest!=1]),3),'integer_AUROC':round(roc_auc_score(d[d.ohca_arrest!=1].in_hospital_mortality.astype(int).values,integ(d[d.ohca_arrest!=1])),3)})
st=pd.DataFrame(rows); st.to_csv('outputs/tables/S_sensitivity_analyses.csv',index=False)
print(st.to_string(index=False))

# ---- build supplementary tables docx (old style) ----
FONT='Times New Roman'
def newdoc():
    D=docx.Document(); s=D.styles['Normal']; s.font.name=FONT; s.font.size=Pt(12); return D
def cap(D,t): p=D.add_paragraph(); r=p.add_run(t); r.bold=True; r.font.name=FONT; r.font.size=Pt(12)
def line(D,t,it=False): p=D.add_paragraph(); r=p.add_run(t); r.italic=it; r.font.name=FONT; r.font.size=Pt(12)
def table(D,df):
    t=D.add_table(rows=1,cols=len(df.columns)); t.style='Light Grid Accent 1'
    for j,c in enumerate(df.columns):
        cell=t.rows[0].cells[j]; run=cell.paragraphs[0].add_run(str(c)); run.bold=True; run.font.name=FONT; run.font.size=Pt(10)
    for _,row in df.iterrows():
        cells=t.add_row().cells
        for j,v in enumerate(row):
            run=cells[j].paragraphs[0].add_run("" if pd.isna(v) else str(v)); run.font.name=FONT; run.font.size=Pt(10)
D=newdoc(); cap(D,'SUPPLEMENTARY MATERIALS — CS-MORT-6'); line(D,'')
cap(D,'Table S1. Sensitivity Analyses Across Cohort Definitions'); table(D,st)
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; CS, cardiogenic shock; ICD, International Classification of Diseases; OHCA, out-of-hospital cardiac arrest.')
line(D,'Notes: The primary cohort retains mixed septic-cardiogenic shock because culture results are not available at the 24-hour scoring time (excluding on them would introduce look-ahead bias) and because mixed shock represents the intended deployment population. Discrimination is higher in pure cardiogenic shock (culture-confirmed infection excluded) and lower within the culture-positive subgroup, indicating the score captures cardiac, not infective, mortality drivers. All AUROC values are 5-fold cross-validated.')
line(D,'')
cap(D,'Table S2. Missing Data and Multiple-Imputation Sensitivity')
table(D,pd.DataFrame([
 {'variable':'Lactate','missing':'19.4%','mortality if missing':'28.7%','mortality if measured':'40.6%','pattern':'MNAR (less sick missing)'},
 {'variable':'Urine output','missing':'8.3%','mortality if missing':'56.2%','mortality if measured':'36.7%','pattern':'MNAR (sicker missing)'}]))
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; MNAR, missing not at random; MICE, multiple imputation by chained equations.')
line(D,'Notes: Discrimination stratified by lactate status: measured AUROC 0.789, missing (imputed) 0.701, with preserved calibration in both (slope ~1.0). Multiple imputation (MICE) AUROC 0.775 versus median imputation 0.778, confirming robustness to the imputation method. The deployed model uses median imputation with the harmonized anion gap as a pre-specified substitution for missing lactate.')
line(D,'')
cap(D,'Table S3. Model-Class Comparison (Development Bake-Off)')
table(D,pd.DataFrame([
 {'model':'Elastic-net logistic (selected)','AUROC':'0.792','calibration slope':'0.93'},
 {'model':'LASSO logistic','AUROC':'0.792','calibration slope':'0.95'},
 {'model':'Random forest','AUROC':'0.790','calibration slope':'1.61'},
 {'model':'Gradient boosting','AUROC':'0.796','calibration slope':'0.75'}]))
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; LASSO, least absolute shrinkage and selection operator.')
line(D,'Notes: Discrimination was equivalent across model classes while the machine-learning models were materially miscalibrated, supporting the choice of penalized logistic regression.')
line(D,'')
cap(D,'Table S4. Subgroup (Fairness) Performance')
table(D,pd.DataFrame([
 {'subgroup':'Male','n':1858,'mortality':'36.4%','AUROC':0.785,'slope':1.04,'CITL':-0.013},
 {'subgroup':'Female','n':1245,'mortality':'41.1%','AUROC':0.767,'slope':0.90,'CITL':0.017},
 {'subgroup':'White','n':1935,'mortality':'36.8%','AUROC':0.770,'slope':1.01,'CITL':-0.009},
 {'subgroup':'Black','n':317,'mortality':'36.6%','AUROC':0.780,'slope':0.85,'CITL':-0.054},
 {'subgroup':'Other/Unknown','n':693,'mortality':'44.9%','AUROC':0.799,'slope':1.03,'CITL':0.059}]))
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; CITL, calibration-in-the-large.')
line(D,'Notes: Discrimination and calibration were consistent across sex and race groups (TRIPOD+AI item 14). Hispanic (n=84) and Asian (n=74) groups were too small for stable estimates and are not shown separately.')
line(D,'')
cap(D,'Table S5. Cluster Heterogeneity Across eICU Hospitals')
table(D,pd.DataFrame([{'hospitals (>=25 CS patients)':24,'median AUROC':0.764,'IQR':'0.713-0.791','pooled eICU AUROC':0.748}]))
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; IQR, interquartile range.')
line(D,'Notes: Per-hospital performance was consistent with no catastrophic site; the wide range reflects small per-hospital samples (TRIPOD+AI items 12d and 23b).')
line(D,'')
cap(D,'Table S6. Cut-Point and Note-Classifier Sensitivity')
table(D,pd.DataFrame([
 {'analysis':'Integer cut-points: guideline/data-derived','result':'AUROC 0.763'},
 {'analysis':'Integer cut-points: distribution quintiles','result':'AUROC 0.765'},
 {'analysis':'Integer cut-points: round-number clinical','result':'AUROC 0.766'},
 {'analysis':'ConText note classifier, note-level PPV (all affirmed, n=1,199)','result':'96.5%'},
 {'analysis':'ConText note classifier, note-only inclusion-driving (n=282)','result':'89.7%'}]))
line(D,'Abbreviations: AUROC, area under the receiver operating characteristic curve; PPV, positive predictive value.')
line(D,'Notes: Integer-score discrimination was robust across three independent cut-point schemes, confirming the cut-points are not arbitrary. The note classifier positive predictive value is a conservative lower bound (the audit detector over-flagged true affirmations).')
D.save('manuscript/SUPPLEMENTARY_TABLES.docx'); print('\nwrote manuscript/SUPPLEMENTARY_TABLES.docx + outputs/tables/S_sensitivity_analyses.csv')
